"""
POC test script for Resume Screener core engine.
Tests (all in one):
  1. PDF/DOCX parsing (generate sample files, extract text)
  2. Content-hash caching (knowledge repository: no redundant re-parsing)
  3. TF-IDF quick match (score, matched/missing skills, top terms)
  4. Deep semantic match with sentence-transformers (similarity + aligned snippets)
  5. Feedback training loop (user corrections change future scores)
"""
import hashlib
import json
import re
import time
import os

RESULTS = {}

SAMPLE_RESUME = """
John Doe
Senior Software Engineer
Email: john.doe@example.com | Phone: 555-1234

SUMMARY
Experienced software engineer with 6 years building scalable web applications
using Python, FastAPI, and React. Strong background in cloud deployment (AWS),
CI/CD pipelines, and MongoDB database design. Led a team of 4 engineers.

SKILLS
Python, JavaScript, React, FastAPI, MongoDB, PostgreSQL, Docker, AWS, Git,
REST APIs, unit testing, Agile, CI/CD, Redis

EXPERIENCE
Senior Software Engineer - TechCorp (2021-Present)
- Built microservices with FastAPI handling 1M requests/day
- Designed MongoDB schemas and optimized queries reducing latency by 40%
- Mentored junior developers and led code reviews

Software Engineer - WebSolutions (2018-2021)
- Developed React dashboards used by 10k+ users
- Implemented CI/CD with GitHub Actions and Docker

EDUCATION
B.S. Computer Science, State University, 2018
"""

SAMPLE_JD = """
We are hiring a Senior Backend Engineer.

Requirements:
- 5+ years of experience with Python web frameworks (FastAPI or Django)
- Strong experience with MongoDB and Redis
- Experience deploying on AWS with Docker and Kubernetes
- Familiarity with CI/CD pipelines and automated testing
- Experience with GraphQL APIs is a plus
- Leadership or mentoring experience preferred
- Knowledge of Terraform is a bonus

Responsibilities:
- Design and build scalable microservices
- Optimize database performance
- Collaborate with frontend teams using React
"""


def make_sample_files(tmpdir="/tmp/poc_files"):
    os.makedirs(tmpdir, exist_ok=True)
    # DOCX
    from docx import Document
    doc = Document()
    for line in SAMPLE_RESUME.strip().split("\n"):
        doc.add_paragraph(line)
    docx_path = os.path.join(tmpdir, "resume.docx")
    doc.save(docx_path)
    # PDF
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    pdf_path = os.path.join(tmpdir, "resume.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    y = 750
    for line in SAMPLE_RESUME.strip().split("\n"):
        c.drawString(50, y, line[:100])
        y -= 14
        if y < 50:
            c.showPage()
            y = 750
    c.save()
    return pdf_path, docx_path


# ---------------- 1. PARSING ----------------
def parse_pdf(path):
    import pdfplumber
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text.append(t)
    return "\n".join(text)


def parse_docx(path):
    from docx import Document
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def test_parsing():
    pdf_path, docx_path = make_sample_files()
    pdf_text = parse_pdf(pdf_path)
    docx_text = parse_docx(docx_path)
    assert "FastAPI" in pdf_text and "MongoDB" in pdf_text, "PDF parse failed"
    assert "FastAPI" in docx_text and "MongoDB" in docx_text, "DOCX parse failed"
    RESULTS["parsing"] = {"pdf_chars": len(pdf_text), "docx_chars": len(docx_text)}
    return pdf_text, docx_text


# ---------------- 2. CACHE ----------------
CACHE = {}  # simulates Mongo collection keyed by sha256


def content_hash(text):
    return hashlib.sha256(text.strip().lower().encode()).hexdigest()


def cached_parse(raw_text):
    h = content_hash(raw_text)
    if h in CACHE:
        return CACHE[h], True  # cache hit
    parsed = {"text": raw_text, "features": extract_skills(raw_text, BASE_SKILLS)}
    CACHE[h] = parsed
    return parsed, False


def test_cache():
    t0 = time.time()
    _, hit1 = cached_parse(SAMPLE_RESUME)
    first = time.time() - t0
    t0 = time.time()
    _, hit2 = cached_parse(SAMPLE_RESUME)
    second = time.time() - t0
    assert hit1 is False and hit2 is True, "Cache behavior wrong"
    RESULTS["cache"] = {"first_parse_s": round(first, 4), "cached_s": round(second, 6), "hit_on_second": hit2}


# ---------------- 3. SKILLS + TF-IDF ----------------
BASE_SKILLS = [
    "python", "javascript", "typescript", "react", "angular", "vue", "fastapi",
    "django", "flask", "node.js", "mongodb", "postgresql", "mysql", "redis",
    "docker", "kubernetes", "aws", "azure", "gcp", "git", "ci/cd", "graphql",
    "rest apis", "terraform", "agile", "unit testing", "machine learning",
    "java", "c++", "sql", "linux", "microservices", "leadership", "mentoring",
]

# feedback-adjustable skill weights (the "trainable" part)
SKILL_WEIGHTS = {s: 1.0 for s in BASE_SKILLS}
SYNONYMS = {"github actions": "ci/cd", "mentored": "mentoring", "led a team": "leadership"}


def normalize(text):
    text = text.lower()
    text = re.sub(r"[^\w\s\+\#\./-]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_skills(text, skills):
    norm = normalize(text)
    found = set()
    for s in skills:
        pattern = r"(?<![\w])" + re.escape(s) + r"(?![\w])"
        if re.search(pattern, norm):
            found.add(s)
    for syn, canonical in SYNONYMS.items():
        if syn in norm:
            found.add(canonical)
    return sorted(found)


def tfidf_match(resume_text, jd_text):
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity

    vec = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), sublinear_tf=True)
    m = vec.fit_transform([normalize(resume_text), normalize(jd_text)])
    cos = float(cosine_similarity(m[0], m[1])[0][0])

    resume_skills = set(extract_skills(resume_text, BASE_SKILLS))
    jd_skills = set(extract_skills(jd_text, BASE_SKILLS))
    matched = resume_skills & jd_skills
    missing = jd_skills - resume_skills

    # weighted skill coverage (feedback-adjustable)
    if jd_skills:
        total_w = sum(SKILL_WEIGHTS.get(s, 1.0) for s in jd_skills)
        matched_w = sum(SKILL_WEIGHTS.get(s, 1.0) for s in matched)
        skill_cov = matched_w / total_w
    else:
        skill_cov = 0.0

    score = round((0.45 * skill_cov + 0.35 * min(cos * 2.5, 1.0) + 0.20 * min(len(matched) / 10, 1.0)) * 100, 1)
    return {
        "score": score,
        "cosine": round(cos, 4),
        "matched_skills": sorted(matched),
        "missing_skills": sorted(missing),
        "skill_coverage": round(skill_cov * 100, 1),
    }


def test_tfidf():
    t0 = time.time()
    res = tfidf_match(SAMPLE_RESUME, SAMPLE_JD)
    dt = time.time() - t0
    assert 30 <= res["score"] <= 100, f"Suspicious score {res['score']}"
    assert "python" in res["matched_skills"] and "mongodb" in res["matched_skills"]
    assert "kubernetes" in res["missing_skills"] and "graphql" in res["missing_skills"]
    # sanity: unrelated JD scores lower
    unrelated = tfidf_match(SAMPLE_RESUME, "We need a pastry chef with baking, cake decorating and food safety skills.")
    assert unrelated["score"] < res["score"], "Unrelated JD should score lower"
    RESULTS["tfidf"] = {**res, "time_s": round(dt, 3), "unrelated_jd_score": unrelated["score"]}
    return res


# ---------------- 4. DEEP EMBEDDINGS ----------------
def test_embeddings():
    from sentence_transformers import SentenceTransformer, util
    t0 = time.time()
    model = SentenceTransformer("all-MiniLM-L6-v2")
    load_t = time.time() - t0

    t0 = time.time()
    doc_emb = model.encode([SAMPLE_RESUME, SAMPLE_JD], normalize_embeddings=True)
    overall = float(util.cos_sim(doc_emb[0], doc_emb[1])[0][0])

    # sentence-level alignment: for each JD requirement find best resume line
    jd_lines = [l.strip("- ").strip() for l in SAMPLE_JD.split("\n") if len(l.strip()) > 20]
    resume_lines = [l.strip("- ").strip() for l in SAMPLE_RESUME.split("\n") if len(l.strip()) > 20]
    jd_emb = model.encode(jd_lines, normalize_embeddings=True)
    r_emb = model.encode(resume_lines, normalize_embeddings=True)
    sims = util.cos_sim(jd_emb, r_emb)
    alignments = []
    for i, jl in enumerate(jd_lines):
        best_j = int(sims[i].argmax())
        alignments.append({"jd": jl[:60], "best_resume_line": resume_lines[best_j][:60], "sim": round(float(sims[i][best_j]), 3)})
    infer_t = time.time() - t0

    assert overall > 0.3, f"Semantic sim too low: {overall}"
    strong = [a for a in alignments if a["sim"] > 0.5]
    assert len(strong) >= 3, "Expected at least 3 strong alignments"
    RESULTS["embeddings"] = {
        "overall_similarity": round(overall, 4),
        "model_load_s": round(load_t, 2),
        "inference_s": round(infer_t, 2),
        "sample_alignments": alignments[:4],
    }


# ---------------- 5. FEEDBACK TRAINING ----------------
def apply_feedback(feedback):
    """feedback: {skill: 'correct'|'incorrect'|'important'} adjusts weights; also learns synonyms."""
    for skill, verdict in feedback.get("skill_feedback", {}).items():
        s = skill.lower()
        if s not in SKILL_WEIGHTS:
            SKILL_WEIGHTS[s] = 1.0
            BASE_SKILLS.append(s)
        if verdict == "important":
            SKILL_WEIGHTS[s] = min(SKILL_WEIGHTS[s] + 0.5, 3.0)
        elif verdict == "incorrect":
            SKILL_WEIGHTS[s] = max(SKILL_WEIGHTS[s] - 0.5, 0.1)
    for syn, canonical in feedback.get("new_synonyms", {}).items():
        SYNONYMS[syn.lower()] = canonical.lower()


def test_feedback():
    before = tfidf_match(SAMPLE_RESUME, SAMPLE_JD)
    # user says kubernetes & terraform are not that relevant (they were 'nice to have'),
    # and that "GraphQL" importance should increase; also adds a new skill via feedback
    apply_feedback({
        "skill_feedback": {"kubernetes": "incorrect", "terraform": "incorrect", "mongodb": "important"},
        "new_synonyms": {"k8s": "kubernetes"},
    })
    after = tfidf_match(SAMPLE_RESUME, SAMPLE_JD)
    assert after["score"] > before["score"], f"Feedback should raise score ({before['score']} -> {after['score']})"
    assert SYNONYMS.get("k8s") == "kubernetes"
    RESULTS["feedback"] = {"score_before": before["score"], "score_after": after["score"],
                           "weights_sample": {k: SKILL_WEIGHTS[k] for k in ["kubernetes", "terraform", "mongodb"]}}


if __name__ == "__main__":
    failures = []
    for name, fn in [("parsing", test_parsing), ("cache", test_cache), ("tfidf", test_tfidf),
                     ("embeddings", test_embeddings), ("feedback", test_feedback)]:
        try:
            fn()
            print(f"[PASS] {name}")
        except Exception as e:
            failures.append((name, str(e)))
            print(f"[FAIL] {name}: {e}")
    print("\n=== RESULTS ===")
    print(json.dumps(RESULTS, indent=2))
    if failures:
        print("\nFAILURES:", failures)
        raise SystemExit(1)
    print("\nALL CORE POC TESTS PASSED")
