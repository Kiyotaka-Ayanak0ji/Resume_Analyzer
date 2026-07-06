"""Resume file parsing (PDF/DOCX/TXT) + content hashing for the knowledge cache."""
import hashlib
import io
import re

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class ParseError(Exception):
    pass


def content_hash(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text.strip().lower())
    return hashlib.sha256(normalized.encode()).hexdigest()


def parse_pdf_bytes(data: bytes) -> str:
    import pdfplumber
    try:
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        text = "\n".join(pages).strip()
        if not text:
            raise ParseError("Could not read any text from this PDF. It may be a scanned image - please paste your resume text instead.")
        return text
    except ParseError:
        raise
    except Exception:
        raise ParseError("This PDF could not be opened. Please check the file or paste your resume text instead.")


def parse_docx_bytes(data: bytes) -> str:
    from docx import Document
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = "\n".join(parts).strip()
        if not text:
            raise ParseError("This document appears to be empty.")
        return text
    except ParseError:
        raise
    except Exception:
        raise ParseError("This DOCX file could not be opened. Please check the file or paste your resume text instead.")


def parse_file(filename: str, data: bytes) -> str:
    if len(data) > MAX_FILE_SIZE:
        raise ParseError("File is too large. Maximum size is 5 MB.")
    if len(data) == 0:
        raise ParseError("The uploaded file is empty.")
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        return parse_pdf_bytes(data)
    if lower.endswith(".docx"):
        return parse_docx_bytes(data)
    if lower.endswith(".txt"):
        try:
            return data.decode("utf-8", errors="replace").strip()
        except Exception:
            raise ParseError("Could not read this text file.")
    raise ParseError("Unsupported file type. Please upload a PDF, DOCX or TXT file.")
